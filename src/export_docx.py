from docx import Document

def export_to_docx(questions, answers_map, out_path):
    doc = Document()
    doc.add_heading("Questionnaire Response", level=1)

    for q in questions:
        doc.add_paragraph(q)

        a = answers_map.get(q, {})
        final_answer = a.get("edited_answer") or a.get("answer") or "Not found in references."
        citations = a.get("citations", [])

        doc.add_paragraph(f"Answer: {final_answer}")
        doc.add_paragraph(f"Citations: {', '.join(citations) if citations else 'None'}")
        doc.add_paragraph("")

    doc.save(out_path)